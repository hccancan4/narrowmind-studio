"""DOCX extraction via python-docx."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document as DocxDocument


def extract(path: Path) -> tuple[str, str, dict[str, Any]]:
    doc = DocxDocument(str(path))
    # python-docx exposes paragraphs in document order. Headings + body all come through
    # via `text`; we keep them flat and let the chunker decide on splits.
    paragraphs: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    # Pull plain text out of any tables as well — important for technical docs.
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                paragraphs.append(" | ".join(row_cells))

    body = "\n\n".join(paragraphs)
    title = path.stem
    core = doc.core_properties
    if core and core.title:
        title = core.title
    elif paragraphs:
        title = paragraphs[0][:120]

    metadata: dict[str, Any] = {
        "format": "docx",
        "paragraph_count": len(paragraphs),
    }
    if core and core.author:
        metadata["docx_author"] = core.author
    if core and core.created:
        metadata["docx_created"] = core.created.isoformat()

    return title, body, metadata
