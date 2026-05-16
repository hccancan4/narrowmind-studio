"""Shared pytest fixtures. Generates PDF / DOCX binary fixtures at session scope so we
don't commit binaries into the repo."""

from __future__ import annotations

from pathlib import Path

import pytest

FIXTURE_DIR = Path(__file__).parent / "fixtures"


@pytest.fixture(scope="session")
def pdf_fixture() -> Path:
    """Generate a tiny one-page PDF the PyMuPDF handler can read."""

    import fitz  # pymupdf

    path = FIXTURE_DIR / "_generated_sample.pdf"
    if path.exists():
        return path
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        "The hard problem of consciousness\n\nThis is a generated PDF fixture for ingestion tests.\nSecond paragraph for chunkability.",
    )
    doc.set_metadata({"title": "Hard problem (test)", "author": "fixture"})
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture(scope="session")
def docx_fixture() -> Path:
    """Generate a small DOCX the python-docx handler can read."""

    from docx import Document

    path = FIXTURE_DIR / "_generated_sample.docx"
    if path.exists():
        return path
    doc = Document()
    doc.core_properties.title = "Functionalism (test)"
    doc.core_properties.author = "fixture"
    doc.add_heading("Functionalism", level=1)
    doc.add_paragraph(
        "This is a generated DOCX fixture for the ingestion handler test. "
        "It has two paragraphs so the body is non-trivial."
    )
    doc.add_paragraph("Second paragraph mentioning multiple realizability.")
    doc.save(str(path))
    return path
